from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth import login, logout, authenticate, update_session_auth_hash
from .forms import TaskForm, TaskFormPart1, TaskFormPart2, TaskFormPart3, UsuarioForm, CustomPasswordForm, TalentoHumanoForm, TalentoHumanoFormPart1, TalentoHumanoFormPart2, TalentoHumanoFormPart3, TalentoHumanoFormPart4, TalentoHumanoAsignacionForm, RelacionTHCForm, ComentarioForm, UsuarioEditForm, UserEditForm, ArchivoForm, CorrespondeciaForm, AdminEditForm, InventarioForm, ObservacionForm, SedeForm, FotoForm, VisitaForm
from .models import Task, Usuario, TalentoHumano, RelacionTHC, Comentario, Archivo, Correspondencia, Inventario, Observacion, Sede, Foto, Visita
from decimal import Decimal
from datetime import datetime
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from datetime import date
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.contrib.auth.models import User
from django.db.models import Q
import os
from docx import Document
from django.http import FileResponse
from io import BytesIO
from docx.shared import Pt
from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openpyxl import Workbook
from django.http import HttpResponse
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from django.db.models import Max
from django.db import IntegrityError
from openpyxl.styles import NamedStyle, Font
from openpyxl.utils import get_column_letter

@login_required
def home(request): 
    return render(request, 'home.html')

@login_required    
def tasks(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    search_term = request.GET.get('search')
    tasks = Task.objects.filter(datecompleted__isnull=True).order_by('-created')

    if search_term:
        tasks = tasks.filter(
            Q(document__icontains=search_term) | 
            Q(nombre__icontains=search_term)
            )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(tasks, page_size)
    page = request.GET.get('page')

    try: 
        tasks = paginator.page(page)
    except PageNotAnInteger:
        tasks = paginator.page(1)
    except EmptyPage:
        tasks = paginator.page(paginator.num_pages)

    return render(request, 'tasks.html', {
        'tasks':tasks,
        'search_term': search_term,
    })

@login_required 
def tasks_completed(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    search_term = request.GET.get('search')
    tasks = Task.objects.filter(datecompleted__isnull=False).order_by('-created')

    if search_term:
        tasks = tasks.filter(
            Q(document__icontains=search_term) | 
            Q(nombre__icontains=search_term)
            )
        
    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(tasks, page_size)
    page = request.GET.get('page')
    
    try:
        tasks = paginator.page(page)
    except PageNotAnInteger:
        tasks = paginator.page(1)
    except EmptyPage:
        tasks = paginator.page(paginator.num_pages)

    return render(request, 'tasks_completed.html', {
        'tasks':tasks,
        'search_term': search_term,
    })

@login_required
def delete_task(request, task_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType != 'Admin':
        return redirect('tasks')

    task = get_object_or_404(Task, pk=task_id, user=request.user)
    if request.method == 'POST':
        task.delete()
        if task.datecompleted is not None:
            return redirect('tasks_completed')
        else:
            return redirect('tasks')      

@login_required         
def task_detail(request, task_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.method == 'GET':
        task = get_object_or_404(Task, pk=task_id)
        talentohumanos = task.talentohumanos.all()
        archivos = Archivo.objects.filter(task=task)
        form = TaskForm(instance=task)
        return render(request, 'task_detail.html', {
            'task': task,
            'form': form, 
            'talentohumanos': talentohumanos,
            'archivos': archivos
        })
    else:
        if request.user.usuario.userType != 'Admin':
            return redirect('tasks')    
        try:
            task = get_object_or_404(Task, pk=task_id)   
            talentohumanos = task.talentohumanos.all() 
            archivos = Archivo.objects.filter(task=task)
            form = TaskForm(request.POST, request.FILES, instance=task)
            form.save()
            if task.datecompleted is not None:
                return redirect('tasks_completed')
            else:
                return redirect('tasks') 
        except ValueError:
            return render(request, 'task_detail.html', {
            'task': task,
            'form': form,
            'talentohumanos': talentohumanos,
            'archivos': archivos,
            'error': 'Ocurrio un error actualizando este contrato'
        })

@login_required 
def complete_task(request, task_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType != 'Admin':
        return redirect('tasks')
    
    task = get_object_or_404(Task, pk=task_id, user=request.user)
    if request.method == 'POST':
        task.datecompleted = timezone.now()
        task.save()
        return redirect('tasks')

@login_required
def add_archivo(request, task_id):
    
    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    task = get_object_or_404(Task, pk=task_id)

    if request.method == 'POST':
        if request.user.usuario.userType != 'Admin':
            return redirect('tasks')

        form = ArchivoForm(request.POST, request.FILES, task = task)

        if form.is_valid():
            form.save()
            return redirect('task_detail', task_id = task_id)
        
    else:
        form = ArchivoForm(task=task)

    return render(request, 'add_archivo.html', {
        'task': task,
        'form': form
    })

@login_required 
def signout(request):
    logout(request)
    return redirect('signin')

def signin(request):

    if request.user.is_authenticated:
        return redirect('home')
    
    if request.method == 'GET':
        return render(request, 'signin.html', {
            'form': AuthenticationForm
        })
    else:
        user = authenticate(request, username=request.POST['username'], password=request.POST['password'])
        if user is None:
            return render(request, 'signin.html', {
            'form': AuthenticationForm,
            'error': 'Nombre de usuario y/o contraseña incorrectos'
            })
        else:
            login(request, user)
            return redirect('home')       
        
def convert_to_serializable(data):

    for key, value in data.items():
        if isinstance(value, Decimal):
            data[key] = str(value)  
        elif isinstance(value, date):
            data[key] = value.strftime('%Y-%m-%d')  
    return data

@login_required 
def task_wizard(request, step=1):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType != 'Admin':
        return redirect('tasks')
    
    if step == 1:
        form = TaskFormPart1(request.POST or None)
    elif step == 2:
        form = TaskFormPart2(request.POST or None)
    elif step == 3:
        form = TaskFormPart3(request.POST or None, request.FILES or None)
    else: 
        return redirect('tasks')

    if request.method == 'POST' and form.is_valid():
        form_data_serializable = convert_to_serializable(form.cleaned_data)
        request.session['form_data_step_{}'.format(step)] = form_data_serializable
        if step < 3:
            return redirect('task_wizard', step=step + 1)
        else:
            form_data = {}
            for i in range(1, 4):
                form_data.update(request.session.pop('form_data_step_{}'.format(i), {}))
            new_task = Task(**form_data, user=request.user)
            new_task.save()
            return redirect('tasks')
    return render(request, 'task_wizard.html', {
        'form': form,
        'step': step
        })

@login_required
def talentohumano_list(request, task_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    task = Task.objects.get(pk=task_id)
    talentohumanos = TalentoHumano.objects.filter(tasks=task)
    search_term = request.GET.get('search')

    if search_term:
        talentohumanos = talentohumanos.filter(
            Q(document__icontains=search_term) | 
            Q(firstName__icontains=search_term) |
            Q(secondName__icontains=search_term) |
            Q(firstLastname__icontains=search_term) |
            Q(secondLastname__icontains=search_term)  
        )

    return render(request, 'talentohumano_list.html', {
        'talentohumanos': talentohumanos,
        'task': task,
        'search_term': search_term
        })

@login_required
def crear_user(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType != 'Admin':
        return redirect('usuarios')
    
    error = None

    if request.method == 'POST':
        form = UsuarioForm(request.POST)
        if form.is_valid():
            user = User.objects.create_user(
                username=form.cleaned_data['username'],
                password=form.cleaned_data['password1'],
                first_name=form.cleaned_data['first_name'],
                last_name=form.cleaned_data['last_name']
            )

                
            usuario = Usuario(
                user=user,
                document=form.cleaned_data['document'],
                red=form.cleaned_data['red'],
                userType=form.cleaned_data['userType']
            )
            usuario.save()

            return redirect('usuarios')

    else:
        form = UsuarioForm()

    context = {
        'form': form
    }

    return render(request, 'crear_usuario.html', context)
    
@login_required    
def usuarios(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    search_term = request.GET.get('search')
    usuarios = Usuario.objects.all().order_by('user__first_name')
    if search_term:
        usuarios = usuarios.filter(
            Q(document__icontains=search_term) | 
            Q(user__first_name__icontains=search_term) |
            Q(user__last_name__icontains=search_term)
            )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(usuarios, page_size)
    page = request.GET.get('page')

    try: 
        usuarios = paginator.page(page)
    except PageNotAnInteger:
        usuarios = paginator.page(1)
    except EmptyPage:
        usuarios = paginator.page(paginator.num_pages)

    return render(request, 'usuarios.html', {
        'usuarios':usuarios,
        'search_term': search_term,
    })

@login_required
def delete_usuario(request, usuario_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType != 'Admin':
        return redirect('tasks')
    
    usuario = get_object_or_404(Usuario, pk=usuario_id)
    if request.method == 'POST':
        if usuario.user:
            user_id = usuario.user.id
            User.objects.filter(id=user_id).delete()
        usuario.delete()
        return redirect('usuarios')
        

@login_required
def usuario_detail(request, usuario_id): 
    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    usuario = get_object_or_404(Usuario, pk=usuario_id)

    if request.method == 'POST':

        if request.user.usuario.userType == 'Admin':
            admin_form = AdminEditForm(request.POST, instance=usuario)
            if admin_form.is_valid():
                user_instance = usuario.user
                password1 = admin_form.cleaned_data['password1']
                password2 = admin_form.cleaned_data['password2']

                if password1 != password2:
                    error = "Las contraseñas no coinciden. Por favor, inténtelo de nuevo."
                    return render(request, 'usuario_detail.html', {
                        'usuario': usuario,
                        'admin_form': admin_form,
                        'error': error
                    })

                if password1:
                    user_instance.set_password(password1)
                    user_instance.save()
                    update_session_auth_hash(request, user_instance)  # Important to keep the user logged in

                usuario = admin_form.save()
                return redirect('usuarios')

            else:
                error = "Ha ocurrido un error. Por favor, vuelva a intentarlo."
                return render(request, 'usuario_detail.html', {
                    'usuario': usuario,
                    'admin_form': admin_form,
                    'error': error
        })
        else:
            usuario_form = UsuarioEditForm(request.POST, instance=usuario)
            user_form = UserEditForm(request.POST, instance=usuario.user)

            if usuario_form.is_valid() and user_form.is_valid():
                usuario_form.save()
                user_form.save()
                return redirect('usuarios')
            else:
                error = "Ha ocurrido un error. Por favor, vuelva a intentarlo."
                return render(request, 'usuario_detail.html', {
                    'usuario': usuario,
                    'usuario_form': usuario_form,
                    'user_form': user_form,
                    'error': error
                })
    else:
        if request.user.usuario.userType == 'Admin' and request.user.usuario.id == usuario_id:
            usuario_form = UsuarioEditForm(instance=usuario)
            user_form = UserEditForm(instance=usuario.user)
            return render(request, 'usuario_detail.html', {
                'usuario': usuario,
                'usuario_form': usuario_form,
                'user_form': user_form
            })
        elif request.user.usuario.userType == 'Admin':
            admin_form = AdminEditForm(instance=usuario)
            return render(request, 'usuario_detail.html', {
                'usuario': usuario,
                'admin_form': admin_form
            })
        else:
            usuario_form = UsuarioEditForm(instance=usuario)
            user_form = UserEditForm(instance=usuario.user)
            return render(request, 'usuario_detail.html', {
                'usuario': usuario,
                'usuario_form': usuario_form,
                'user_form': user_form
            })

@login_required
def password_change(request, usuario_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    usuario = get_object_or_404(Usuario, pk=usuario_id)

    if request.method == 'POST':
        
        if request.user.usuario.userType != 'Admin':
            return redirect('usuarios')
        
        password_form = CustomPasswordForm(request.user, request.POST)
        if password_form.is_valid():
            user = password_form.save()
            update_session_auth_hash(request, user)
            return redirect('usuarios')
        
    else:
        if request.user.usuario.userType != 'Admin':
            return redirect('usuarios')
        
        password_form = CustomPasswordForm(request.user)

    return render(request, 'usuario_password.html', {
        'usuario': usuario,
        'password_form': password_form
    })

@login_required    
def talentohumanos(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    search_term = request.GET.get('search')
    talentohumanos = TalentoHumano.objects.all().order_by('firstName')
    if search_term:
        talentohumanos = talentohumanos.filter(
            Q(document__icontains=search_term) | 
            Q(firstName__icontains=search_term) |
            Q(secondName__icontains=search_term) |
            Q(firstLastname__icontains=search_term) |
            Q(secondLastname__icontains=search_term)
            )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(talentohumanos, page_size)
    page = request.GET.get('page')
    
    try: 
        talentohumanos = paginator.page(page)
    except PageNotAnInteger:
        talentohumanos = paginator.page(1)
    except EmptyPage:
        talentohumanos = paginator.page(paginator.num_pages)

    return render(request, 'talentohumanos.html', {
        'talentohumanos': talentohumanos,
        'search_term': search_term,
    })

@login_required
def delete_talentohumano(request, talentohumano_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'Auxiliar':
        return redirect('talentohumanos')
    
    talentohumano = get_object_or_404(TalentoHumano, pk=talentohumano_id)
    if request.method == 'POST':
        talentohumano.delete()
        return redirect('talentohumanos')

@login_required
def talentohumano_wizard(request, step=1):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'Auxiliar':
        return redirect('talentohumanos')
    
    if step == 1:
        form = TalentoHumanoFormPart1(request.POST or None)
    elif step == 2:
        form = TalentoHumanoFormPart2(request.POST or None)
    elif step == 3:
        form = TalentoHumanoFormPart3(request.POST or None)
    elif step == 4:
        form = TalentoHumanoFormPart4(request.POST or None, request.FILES or None)
    else: 
        return redirect('talentohumanos')
    
    if request.method == 'POST' and form.is_valid():
        form_data_serializable = convert_to_serializable(form.cleaned_data)
        request.session['form_data_step_{}'.format(step)] = form_data_serializable
        if step < 4:
            return redirect('talentohumano_wizard', step=step + 1)
        else:
            form_data = {}
            for i in range(1, 5):
                form_data.update(request.session.pop('form_data_step_{}'.format(i), {}))
            new_talentohumano = TalentoHumano(**form_data)
            new_talentohumano.save()    
            return redirect('talentohumanos')
    return render(request, 'talentohumano_wizard.html', {
        'form': form,
        'step': step
        })    

@login_required
def add_comment(request, talentohumano_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    talentohumano = get_object_or_404(TalentoHumano, pk = talentohumano_id)

    if request.method == 'POST':
        form = ComentarioForm(request.POST)
        if form.is_valid():
            comentario = form.save(commit=False)
            comentario.talentohumano = talentohumano
            comentario.usuario = Usuario.objects.get(pk=request.user.id)
            comentario.save()
            return redirect('talentohumano_detail', talentohumano_id = talentohumano.id)
    else:
        form = ComentarioForm()

    return render(request, 'add_comment.html', {
        'form': form,
        'talentohumano': talentohumano
    })

@login_required         
def talentohumano_detail(request, talentohumano_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.method == 'GET':
        talentohumano = get_object_or_404(TalentoHumano, pk=talentohumano_id)
        tasks = talentohumano.tasks.all()
        comentarios = Comentario.objects.filter(talentohumano=talentohumano).order_by('-date_made')

        page_size = int(request.GET.get('page_size', 5))
        paginator = Paginator(comentarios, page_size)  
        page = request.GET.get('page')

        try:
            comentarios = paginator.page(page)
        except PageNotAnInteger:
            comentarios = paginator.page(1)
        except EmptyPage:
            comentarios = paginator.page(paginator.num_pages)

        form = TalentoHumanoForm(instance=talentohumano)
        return render(request, 'talentohumano_detail.html', {
            'talentohumano': talentohumano,
            'form': form,
            'tasks': tasks,
            'comentarios': comentarios
        })
    else:
        try:
            if request.user.usuario.userType == 'Auxiliar':
                return redirect('talentohumanos')
            
            talentohumano = get_object_or_404(TalentoHumano, pk=talentohumano_id) 
            tasks = talentohumano.tasks.all()   
            form = TalentoHumanoForm(request.POST, request.FILES, instance=talentohumano)
            form.save()
            return redirect('talentohumanos')
        except ValueError:
            return render(request, 'talentohumano_detail.html', {
            'talentohumano': talentohumano,
            'form': form,
            'tasks': tasks,
            'error': 'Ocurrio un error actualizando este talento humano'
        })

@login_required
def asignacion(request, task_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    if request.user.usuario.userType == 'Auxiliar':
        return redirect('tasks')
    
    task = Task.objects.get(id=task_id)

    if request.method == 'POST':
        form = TalentoHumanoAsignacionForm(request.POST)
        if form.is_valid():
            selected_talentohumanos = form.cleaned_data['talentohumanos']
            for talentohumano in selected_talentohumanos:
                try:
                    relacion = RelacionTHC.objects.create(
                        talentoHumano = talentohumano,
                        task = task,
                        contractType = None,
                        modalidad = None,
                        cargo = None,
                        jornada = None,
                        startDate = None,
                        endDate = None,
                        monthlySalary = None,
                        totalSalary = None,
                        functions = None,
                        contractsObject = None,
                        lugarPrestacionServicios = None,
                        departamentoPrestacionServicios = None,
                        paymentMethod = None
                    )

                except IntegrityError:
                    error = "Un empleado no se puede asignar mas de una vez a el mismo contrato, por favor vuelva a intentarlo"
                    return render(request, 'asignacion.html', {
                        'form': form,
                        'task': task,
                        'error': error
                    })
                
            if task.datecompleted is not None:
                return redirect('tasks_completed')
            else:
                return redirect('tasks')
    else:
        form = TalentoHumanoAsignacionForm()

        return render(request, 'asignacion.html', {
            'form': form,
            'task': task
        })
    
    return HttpResponse("Ha ocurrido un error")
    
@login_required
def relaciones(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    search_term = request.GET.get('search')
    relaciones = RelacionTHC.objects.all().order_by('-created')
                                           
    if search_term:
        relaciones = relaciones.filter(
            Q(task__document__icontains=search_term) | 
            Q(task__nombre__icontains=search_term) |
            Q(talentoHumano__firstName__icontains=search_term) |
            Q(talentoHumano__secondName__icontains=search_term) |
            Q(talentoHumano__firstLastname__icontains=search_term) |
            Q(talentoHumano__secondLastname__icontains=search_term) |
            Q(talentoHumano__document__icontains=search_term) 
            )
        
    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(relaciones, page_size)
    page = request.GET.get('page')

    try: 
        relaciones = paginator.page(page)
    except PageNotAnInteger:
        relaciones = paginator.page(1)
    except EmptyPage:
        relaciones = paginator.page(paginator.num_pages)

    return render(request, 'relaciones.html', {
        'relaciones': relaciones,
        'search_term': search_term,
    })

@login_required
def delete_relacion(request, relacion_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'Auxiliar':
        return redirect('relaciones')
    
    relacion = get_object_or_404(RelacionTHC, pk=relacion_id)
    if request.method == 'POST':
        relacion.delete()
        return redirect('relaciones')

@login_required         
def relacion_detail(request, relacion_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    relacion = get_object_or_404(RelacionTHC, pk=relacion_id)
    if request.method == 'POST':
        if request.user.usuario.userType == 'Auxiliar':
            return redirect('relaciones')
        
        form = RelacionTHCForm(request.POST, instance=relacion)
        if form.is_valid():
            form.save()
            return redirect('relaciones')
    else:
        form = RelacionTHCForm(instance=relacion)

    return render(request, 'relacion_detail.html', {
        'relacion': relacion,
        'form': form
    })

@login_required
def generarLaboral(request, talentohumano_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    template_path = os.path.join('tasks', 'templates', 'documents', 'Certificado Laboral.docx')
    doc = Document(template_path)

    talentohumano = get_object_or_404(TalentoHumano, id=talentohumano_id)

    placeholders = {
        'nombreCompleto': talentohumano.nombreCompleto,
        'document': talentohumano.document,
        'expeditionPlace': talentohumano.expeditionPlace,
        'contractType': talentohumano.contractType,
        'monthlySalary': str(talentohumano.monthlySalary),
        'cargo': talentohumano.cargo,
        'startDate': str(talentohumano.startDate.strftime('%d-%m-%Y')),
        'endDate': str(talentohumano.endDate.strftime('%d-%m-%Y')),
        'today': datetime.now().strftime('%d-%m-%Y')
    }

    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    response = FileResponse(output)
    response['Content-Disposition'] = f'attachment; filename="Certificado_Laboral_{talentohumano.nombreCompleto}.docx"'

    return response

@login_required
def generarOtro(request, relacion_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    template_path = os.path.join('tasks', 'templates', 'documents', 'Certificado Experiencia.docx')
    doc = Document(template_path)

    relacion = get_object_or_404(RelacionTHC, id=relacion_id)

    if relacion.contractType is None:
        return redirect('correspondencias')

    placeholders = {
        'nombreCompleto': relacion.talentoHumano.nombreCompleto,
        'document': relacion.talentoHumano.document,
        'expeditionPlace': relacion.talentoHumano.expeditionPlace,
        'contractType': relacion.contractType,
        'contractDocument': relacion.task.document,
        'contractor': relacion.task.contractor,
        'contractsObject': relacion.contractsObject,
        'lugarPrestacionServicios': relacion.lugarPrestacionServicios,
        'functions': relacion.functions,
        'monthlySalary': str(relacion.monthlySalary),
        'totalSalary': str(relacion.totalSalary),
        'startDate': str(relacion.startDate.strftime('%d-%m-%Y')),
        'endDate': str(relacion.endDate.strftime('%d-%m-%Y')),
        'today': datetime.now().strftime('%d-%m-%Y')
    }

    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Avenir Light'
                                run.font.size = Pt(6.5)
                                run.font.color.rgb = RGBColor(0, 0, 0) 
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    response = FileResponse(output)
    response['Content-Disposition'] = f'attachment; filename="Certificado_Experiencia_{relacion.talentoHumano.nombreCompleto}.docx"'

    return response

@login_required
def generate_excel_report_tasks(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    if request.user.usuario.userType == 'Auxiliar':
        return redirect('tasks')

    tasks = Task.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Documento", "Año", "Tipo de Contrato", "Nombre de Contrato", "Salario Minimo", "Valor Contrato en SMMLV", "Aporte Contratante", "Aporte ONG La Red", "Valor Total", "Fecha de Inicio", "Fecha de Terminación", "Duración (Meses)", "Duración (Dias)", "Secop", "Rup", "Contratante", "Departamentos", "Municipios", "Tema", "Objeto"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)


    for task in tasks:
        row = [task.document, task.year, task.contractType, task.nombre, task.minimumSalary, task.contractValueSMMLV, task.contractorsBudget, task.redsBudget, task.totalValue, task.startDate.strftime('%d-%m-%Y'), task.endDate.strftime('%d-%m-%Y'), task.durationMonths, task.durationDays, task.secop, task.rup, task.contractor, task.departments, task.municipalities, task.topic, task.object]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    currency_style = NamedStyle(name='currency', number_format='_("$"* #,##0.00_);_("$"* \(#,##0.00\);_("$"* "-"_);_(@_)')

    currency_columns = ["Salario Minimo", "Valor Contrato en SMMLV", "Aporte Contratante", "Aporte ONG La Red", "Valor Total"]
    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1  
        col_letter = get_column_letter(col_index)
        
        for cell in ws[col_letter]:
            cell.style = currency_style
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1
        col_letter = get_column_letter(col_index)
        header_cell = ws[f"{col_letter}1"]

        header_cell.font = Font(bold=True)
        header_cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        header_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        header_cell.border = thin_border
        ws.column_dimensions[col_letter].width = max(len(header_cell.value), 15)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_contratos.xlsx'

    wb.save(response)

    return response

@login_required
def generate__excel_report_talentohumanos(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    if request.user.usuario.userType == 'Auxiliar':
        return redirect('tasks')
    
    talentohumanos = TalentoHumano.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Documento", "Activo", "Tipo de Identificación", "Fecha de Expedición del Documento", "Lugar de Expedición del Documento", "Nombre Completo", "Sexo", "Fecha de Nacimiento", "Lugar de Nacimiento", "Dirección", "Municipio de Residencia", "Telefono", "Telefono del Contacto", "Correo Personal", "Correo Institucional", "Escolaridad", "Profesión", "Estado Civil", "Grupo Sanguineo", "EPS", "AFS", "Codigo del Centro de Trabajo", "Codigo de la Actividad a Desarrollar", "Tipo de Contrato", "Modalidad de Trabajo", "Cargo", "Jornada Laboral", "Fecha de Inicio", "Fecha de Terminación", "Salario Mensual", "Salario Total", "Funciones Especificas", "Objeto del Contrato", "Lugar de Prestación de Servicios", "Departamento de Prestación de Servicios", "Metodo de Pago"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)

    for talentohumano in talentohumanos:
        row = [talentohumano.document, talentohumano.activo, talentohumano.idType, talentohumano.docExpeditionDate, talentohumano.expeditionPlace, talentohumano.nombreCompleto, talentohumano.sex, talentohumano.birthday.strftime('%d-%m-%Y'), talentohumano.birthPlace, talentohumano.residenceAddress, talentohumano.municipalitieResidence, talentohumano.phone, talentohumano.contactPhone, talentohumano.personalEmail, talentohumano.professionalEmail, talentohumano.escolaridad, talentohumano.profesion, talentohumano.estadoCivil, talentohumano.bloodGroup, talentohumano.eps, talentohumano.afs, talentohumano.codigoCentroTrabajo, talentohumano.codigoActividadDesarrollar, talentohumano.contractType, talentohumano.modalidad, talentohumano.cargo, talentohumano.jornada, talentohumano.startDate.strftime('%d-%m-%Y'), talentohumano.endDate.strftime('%d-%m-%Y'), talentohumano.monthlySalary, talentohumano.totalSalary, talentohumano.functions, talentohumano.contractsObject, talentohumano.lugarPrestacionServicios, talentohumano.departamentoPrestacionServicios, talentohumano.paymentMethod]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    currency_style = NamedStyle(name='currency', number_format='_("$"* #,##0.00_);_("$"* \(#,##0.00\);_("$"* "-"_);_(@_)')

    currency_columns = ["Salario Mensual", "Salario Total"]
    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1  
        col_letter = get_column_letter(col_index)
        
        for cell in ws[col_letter]:
            cell.style = currency_style
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1
        col_letter = get_column_letter(col_index)
        header_cell = ws[f"{col_letter}1"]

        header_cell.font = Font(bold=True)
        header_cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        header_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        header_cell.border = thin_border
        ws.column_dimensions[col_letter].width = max(len(header_cell.value), 15)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_talentohumano.xlsx'

    wb.save(response)

    return response

@login_required
def correspondencias(request):
    return render(request, 'correspondencias.html')

@login_required
def correspondencias_int(request):
    
    search_term = request.GET.get('search')
    correspondecias = Correspondencia.objects.all().order_by('dateCreated')

    if search_term:
        correspondecias = correspondecias.filter(
            Q(codigo__icontains=search_term) | 
            Q(año__icontains=search_term) |
            Q(user__username__icontains=search_term) |
            Q(asunto__icontains=search_term) |
            Q(areaProject__icontains=search_term) |
            Q(sentTo__icontains=search_term)
            )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(correspondecias, page_size)
    page = request.GET.get('page')

    
    try: 
        correspondecias = paginator.page(page)
    except PageNotAnInteger:
        correspondecias = paginator.page(1)
    except EmptyPage:
        correspondecias = paginator.page(paginator.num_pages)

    return render(request, 'correspondencias_int.html', {
        'correspondecias':correspondecias,
        'search_term': search_term
    })

@login_required
def correspondencias_ext(request):
    
    search_term = request.GET.get('search')
    correspondecias = Correspondencia.objects.all().order_by('dateCreated')

    if search_term:
        correspondecias = correspondecias.filter(
            Q(codigo__icontains=search_term) | 
            Q(año__icontains=search_term) |
            Q(user__username__icontains=search_term) |
            Q(asunto__icontains=search_term) |
            Q(areaProject__icontains=search_term) |
            Q(sentTo__icontains=search_term)
            )

    paginator = Paginator(correspondecias, 10)
    page = request.GET.get('page')
    
    try: 
        correspondecias = paginator.page(page)
    except PageNotAnInteger:
        correspondecias = paginator.page(1)
    except EmptyPage:
        correspondecias = paginator.page(paginator.num_pages)

    return render(request, 'correspondencias_ext.html', {
        'correspondecias':correspondecias,
        'search_term': search_term
    })

@login_required
def crear_correspondencia_int(request):
   
    if request.method == 'POST':
        form = CorrespondeciaForm(request.POST, request.FILES)
        if form.is_valid():
            correspondencia = form.save(commit=False)
            correspondencia.user = request.user
            correspondencia.tipo = 'Interna'

            tipo = 'Interna'
            año = date.today().year
            last_sequence_number = Correspondencia.objects.filter(tipo=tipo, año=año).aggregate(Max('sequence_number_int'))['sequence_number_int__max']

            if last_sequence_number is not None:
                sequence_number = last_sequence_number + 1
            else:
                sequence_number = 1

            codigo = f"CI-{año}-{sequence_number:04d}"
            correspondencia.codigo = codigo
            correspondencia.sequence_number_int = sequence_number
            correspondencia.save()

                
            return redirect('correspondencias_int')
                
    else:

        form = CorrespondeciaForm()

    return render(request, 'create_correspondencia.html', {
        'form': form
    })

    
@login_required
def crear_correspondencia_ext(request):
   
    if request.method == 'POST':
        form = CorrespondeciaForm(request.POST, request.FILES)
        if form.is_valid():
            correspondencia = form.save(commit=False)
            correspondencia.user = request.user
            correspondencia.tipo = 'Externa'

            tipo = 'Externa'
            año = date.today().year
            last_sequence_number = Correspondencia.objects.filter(tipo=tipo, año=año).aggregate(Max('sequence_number_int'))['sequence_number_int__max']

            if last_sequence_number is not None:
                sequence_number = last_sequence_number + 1
            else:
                sequence_number = 1

            codigo = f"CE-{año}-{sequence_number:04d}"
            correspondencia.codigo = codigo
            correspondencia.sequence_number_int = sequence_number
            correspondencia.save()

                
            return redirect('correspondencias_ext')
                
    else:

        form = CorrespondeciaForm()

    return render(request, 'create_correspondencia.html', {
        'form': form
    })

@login_required         
def correspondencia_detail(request, correspondencia_id):
    correspondencia = get_object_or_404(Correspondencia, pk=correspondencia_id)
    if request.method == 'POST':

        if request.user != correspondencia.user:
            if correspondencia.tipo == 'Externa':

                return redirect('correspondencias_ext')
            
            else:

                return redirect('correspondencias_int')
        
        form = CorrespondeciaForm(request.POST, request.FILES, instance=correspondencia)
        if form.is_valid():
            form.save()

            if correspondencia.tipo == 'Externa':

                return redirect('correspondencias_ext')
            
            else:

                return redirect('correspondencias_int')

    else:
        form = CorrespondeciaForm(instance=correspondencia)

    return render(request, 'correspondencia_detail.html', {
        'correspondencia': correspondencia,
        'form': form
    })

@login_required
def delete_correspondencia(request, correspondencia_id):
    correspondencia = get_object_or_404(Correspondencia, pk=correspondencia_id)
    if request.user != correspondencia.user:

        if correspondencia.tipo == 'Externa':

            return redirect('correspondencias_ext')
        
        else:

            return redirect('correspondencias_int')
        
    if request.method == 'POST':
        correspondencia.delete()
        if correspondencia.tipo == 'Externa':

            return redirect('correspondencias_ext')
        
        else:

            return redirect('correspondencias_int')
    
@login_required
def generate__excel_report_correspondencias(request):

    if request.user.usuario.userType == 'Auxiliar' or request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')

    correspondencias = Correspondencia.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Año", "Codigo", "Tipo", "Fecha de Creación", "Enviado A", "Area o Proyecto", "Asunto", "User"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)

    for correspondencia in correspondencias:
        row = [correspondencia.año, correspondencia.codigo, correspondencia.tipo, correspondencia.dateCreated, correspondencia.sentTo, correspondencia.areaProject, correspondencia.asunto, correspondencia.user.username]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_correspondencias.xlsx'

    wb.save(response)

    return response

@login_required    
def inventarios(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    return render(request, 'inventarios.html')

@login_required
def crear_inventario(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('inventarios')
    
    if request.method == 'POST':
        form = InventarioForm(request.POST)
        if form.is_valid():
            inventario = form.save(commit=False)
            inventario.prestado = False  
            inventario.save()
            return redirect('inventarios')
    
    else:
        form = InventarioForm()
    return render(request, 'crear_inventario.html', {
        'form': form
    })

@login_required
def inventario_total(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    search_term = request.GET.get('search')
    inventarios = Inventario.objects.all().order_by('dateAdquired')

    if search_term:
        inventarios = inventarios.filter(
            Q(codigo__icontains=search_term) | 
            Q(marca__icontains=search_term) | 
            Q(entregadoA__icontains=search_term)
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(inventarios, page_size)
    page = request.GET.get('page')
    
    try: 
        inventarios = paginator.page(page)
    except PageNotAnInteger:
        inventarios = paginator.page(1)
    except EmptyPage:
        inventarios = paginator.page(paginator.num_pages)

    return render(request, 'inventario_total.html', {
        'inventarios':inventarios,
        'search_term': search_term
    })

@login_required
def inventario_disponible(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    search_term = request.GET.get('search')
    inventarios = Inventario.objects.filter(prestado=False).order_by('dateAdquired')

    if search_term:
        inventarios = inventarios.filter(
            Q(codigo__icontains=search_term) | 
            Q(marca__icontains=search_term) | 
            Q(entregadoA__icontains=search_term)
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(inventarios, page_size)
    page = request.GET.get('page')
    
    try: 
        inventarios = paginator.page(page)
    except PageNotAnInteger:
        inventarios = paginator.page(1)
    except EmptyPage:
        inventarios = paginator.page(paginator.num_pages)

    return render(request, 'inventario_disponible.html', {
        'inventarios': inventarios,
        'search_term': search_term
    })

@login_required
def inventario_pendiente(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    search_term = request.GET.get('search')
    inventarios = Inventario.objects.filter(prestado=True).order_by('dateAdquired')

    if search_term:
        inventarios = inventarios.filter(
            Q(codigo__icontains=search_term) | 
            Q(marca__icontains=search_term) | 
            Q(entregadoA__icontains=search_term)
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(inventarios, page_size)
    page = request.GET.get('page')
    
    try: 
        inventarios = paginator.page(page)
    except PageNotAnInteger:
        inventarios = paginator.page(1)
    except EmptyPage:
        inventarios = paginator.page(paginator.num_pages)

    return render(request, 'inventario_pendiente.html', {
        'inventarios': inventarios,
        'search_term': search_term
    })

@login_required         
def inventario_detail(request, inventario_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    inventario = get_object_or_404(Inventario, pk=inventario_id)
    observaciones = Observacion.objects.filter(inventario=inventario)
    if request.method == 'POST':
        form = InventarioForm(request.POST, instance=inventario)
        if form.is_valid():
            inventario = form.save(commit=False)
            if inventario.entregadoA or inventario.redProject:
                inventario.prestado = True
            else:
                inventario.prestado = False
            if inventario.dateGivenBack is not None:
                observacion = Observacion.objects.create(
                    inventario = inventario,
                    observacion = None,
                    entregadoA = inventario.entregadoA,
                    redProject=inventario.redProject,
                    dateTaken=inventario.dateTaken,
                    dateGivenBack=inventario.dateGivenBack
                )

                inventario.entregadoA = None
                inventario.redProject = None
                inventario.dateTaken = None
                inventario.dateGivenBack = None
                inventario.prestado = False
            inventario.save()
            return redirect('inventarios')

    else:
        form = InventarioForm(instance=inventario)

    return render(request, 'inventario_detail.html', {
        'inventario': inventario,
        'observaciones': observaciones,
        'form': form
    })

@login_required
def delete_inventario(request, inventario_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')
    
    inventario = get_object_or_404(Inventario, pk=inventario_id)
        
    if request.method == 'POST':
        inventario.delete()
        return redirect('inventarios')

@login_required
def generate__excel_report_inventarios(request):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')

    inventarios = Inventario.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Nombre", "Cantidad", "Codigo", "Numero Serial", "Marca", "Estado", "Fecha de Aquisición", "Valor", "Descripción", "Caracteristicas", "Entregado A", "Red o Proyecto", "Fecha Inicial de Prestamo", "Fecha de Devolución"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)

    for inventario in inventarios:
        row = [inventario.name, inventario.cantidad, inventario.codigo, inventario.serial_num, inventario.marca, inventario.estado, inventario.dateAdquired, inventario.valor, inventario.descripcion, inventario.caracteristicas, inventario.entregadoA, inventario.redProject, inventario.dateTaken, inventario.dateGivenBack]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    currency_style = NamedStyle(name='currency', number_format='_("$"* #,##0.00_);_("$"* \(#,##0.00\);_("$"* "-"_);_(@_)')

    currency_columns = ["Valor"]
    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1  
        col_letter = get_column_letter(col_index)
        
        for cell in ws[col_letter]:
            cell.style = currency_style
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    for col_name in currency_columns:
        col_index = headers.index(col_name) + 1
        col_letter = get_column_letter(col_index)
        header_cell = ws[f"{col_letter}1"]

        header_cell.font = Font(bold=True)
        header_cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        header_cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        header_cell.border = thin_border
        ws.column_dimensions[col_letter].width = max(len(header_cell.value), 15)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_inventario.xlsx'

    wb.save(response)

    return response

def view_observacion(request, observacion_id):

    if request.user.usuario.userType == 'Correspondencia':
        return redirect('correspondencias')
    
    if request.user.usuario.userType == 'General':
        return redirect('home')  
    
    observacion = get_object_or_404(Observacion, pk=observacion_id)
    form = ObservacionForm(instance=observacion)

    if request.method == 'POST':
        form = ObservacionForm(request.POST, instance=observacion)
        if form.is_valid():
            form.save()
            return redirect('inventarios')

    return render(request, 'observacion_detail.html', {
        'observacion': observacion,
        'form': form
        })

@login_required    
def inspeccionesLocativas(request):
    
    return render(request, 'inspeccionesLocativas.html')

@login_required
def crear_sede(request):

    if request.user.usuario.userType != 'Correspondencia' and request.user.usuario.userType != 'Admin':
        return redirect('sedes')
    
    if request.method == 'POST':
        form = SedeForm(request.POST)
        if form.is_valid():
            sede = form.save(commit=False) 
            sede.save()
            return redirect('sedes')
    
    else:
        form = SedeForm()
        
    return render(request, 'crear_sede.html', {
        'form': form
    })

@login_required
def delete_sede(request, sede_id):

    if request.user.usuario.userType != 'Correspondencia' and request.user.usuario.userType != 'Admin':
        return redirect('sedes')
    
    sede = get_object_or_404(Sede, pk=sede_id)
        
    if request.method == 'POST':
        sede.delete()
        return redirect('sedes')

@login_required
def sedes(request):
    
    search_term = request.GET.get('search')
    sedes = Sede.objects.all()

    if search_term:
        sedes = sedes.filter(
            Q(nombre__icontains=search_term) | 
            Q(departamento__icontains=search_term) | 
            Q(municipio__icontains=search_term) | 
            Q(contrato__document__icontains=search_term) |  
            Q(contrato__nombre__icontains=search_term)
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(sedes, page_size)
    page = request.GET.get('page')
    
    try: 
        sedes = paginator.page(page)
    except PageNotAnInteger:
        sedes = paginator.page(1)
    except EmptyPage:
        sedes = paginator.page(paginator.num_pages)

    return render(request, 'sedes.html', {
        'sedes': sedes,
        'search_term': search_term
    })

@login_required
def sede_detail(request, sede_id):

    sede = get_object_or_404(Sede, pk=sede_id)
    fotos = Foto.objects.filter(sede=sede)

    if request.method == 'POST':

        form = SedeForm(request.POST, instance=sede)

        if form.is_valid():
            form.save()
            return redirect('sedes')  

    else:  
        form = SedeForm(instance=sede)

    return render(request, 'sede_detail.html', {
        'sede': sede,
        'fotos': fotos,
        'form': form
    })

@login_required
def generate_excel_report_sedes(request):

    sedes = Sede.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Nombre", "Departamento", "Municipio", "Dirección", "Latitud", "Longitud", "Grados ° Latitud*", "Minutos ° Latitud*", "Segundos ° Latitud*", "Grados ° Longitud*", "Minutos ° Longitud*", "Segundos ° Longitud*", "Descripción"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)

    for sede in sedes:
        row = [sede.nombre, sede.departamento, sede.municipio, sede.address, sede.latitud, sede.longitud, sede.lat1, sede.lat2, sede.lat3, sede.lon1, sede.lon2, sede.lon3, sede.descripcion]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_sedes.xlsx'

    wb.save(response)

    return response

@login_required
def add_foto(request, sede_id):

    if request.user.usuario.userType != 'Correspondencia' and request.user.usuario.userType != 'Admin':
        return redirect('sedes')
    
    if request.method == 'POST':
        form = FotoForm(request.POST, request.FILES)
        if form.is_valid():
            foto = form.save(commit=False)
            foto.sede_id = sede_id

            last_foto = Foto.objects.filter(sede_id=sede_id).order_by('-number').first()
            if last_foto:
                foto.number = str(int(last_foto.number) + 1)
            else:
                foto.number = '1'

            foto.save()
            return redirect('sede_detail', sede_id=sede_id)
    else:
        form = FotoForm()

    return render(request, 'add_foto.html', {
        'form': form
        })

@login_required    
def visitas(request):
    
    return render(request, 'visitas.html')

@login_required
def crear_visita(request, sede_id):

    if request.user.usuario.userType != 'Correspondencia' and request.user.usuario.userType != 'Admin':
        return redirect('sedes')
    
    sede = get_object_or_404(Sede, pk=sede_id)

    if request.method == 'POST':
        form = VisitaForm(request.POST)
        if form.is_valid():
            form.instance.departamento = sede.departamento
            form.instance.municipio = sede.municipio
            form.instance.address = sede.address
            form.instance.sede = sede

            usuario_instance = Usuario.objects.get(user=request.user)
            form.instance.usuario = usuario_instance

            if 'NO' in form.cleaned_data.values():
                form.instance.estado = 'Pendiente'
            else:
                form.instance.estado = 'Completada'

            visita = form.save()
            return redirect('sedes')
        
    else:
        form = VisitaForm(initial={
            'sede': sede
        })

    return render(request, 'crear_visita.html', {
        'form': form
    })

@login_required
def visita_detail(request, visita_id):

    visita = get_object_or_404(Visita, pk=visita_id)
    fotos = Foto.objects.filter(visita=visita)

    if request.method == 'POST':

        form = VisitaForm(request.POST, instance=visita)

        if form.is_valid():
            form.save()
            return redirect('visitas')  

    else:  
        form = VisitaForm(instance=visita)

    return render(request, 'visita_detail.html', {
        'visita': visita,
        'fotos': fotos,
        'form': form
    })

@login_required
def visitas_completadas(request):
    
    search_term = request.GET.get('search')
    visitas = Visita.objects.filter(estado='Completada')

    if search_term:
        visitas = visitas.filter( 
            Q(departamento__icontains=search_term) | 
            Q(municipio__icontains=search_term) | 
            Q(usuario__user__username__icontains=search_term) 
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(visitas, page_size)
    page = request.GET.get('page')
    
    try: 
        visitas = paginator.page(page)
    except PageNotAnInteger:
        visitas = paginator.page(1)
    except EmptyPage:
        visitas = paginator.page(paginator.num_pages)

    return render(request, 'visitas_completadas.html', {
        'visitas': visitas,
        'search_term': search_term
    })

@login_required
def visitas_pendientes(request):
    
    search_term = request.GET.get('search')
    visitas = Visita.objects.filter(estado='Pendiente')

    if search_term:
        visitas = visitas.filter( 
            Q(departamento__icontains=search_term) | 
            Q(municipio__icontains=search_term) | 
            Q(usuario__user__username__icontains=search_term) 
        )

    page_size = int(request.GET.get('page_size', 10))
    paginator = Paginator(visitas, page_size)
    page = request.GET.get('page')
    
    try: 
        visitas = paginator.page(page)
    except PageNotAnInteger:
        visitas = paginator.page(1)
    except EmptyPage:
        visitas = paginator.page(paginator.num_pages)

    return render(request, 'visitas_pendientes.html', {
        'visitas': visitas,
        'search_term': search_term
    })

@login_required
def generate_excel_report_inspecciones(request):

    visitas = Visita.objects.all()

    wb = Workbook()
    ws = wb.active

    headers = ["Fecha de la Inspección", "Hora de la Inspección", "Departamento", "Municipio", "Dirección", "Usuario", "Sede", "Estado"]
    ws.append(headers)

    header = ws[1]
    thin_border = Border(left=Side(style='thin'), 
                        right=Side(style='thin'), 
                        top=Side(style='thin'), 
                        bottom=Side(style='thin'))
    
    for cell in header:
        cell.font = Font(bold=True)
        cell.fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        cell.border = thin_border
        ws.column_dimensions[cell.column_letter].width = max(len(cell.value), 15)

    for visita in visitas:
        row = [visita.date, visita.hora, visita.departamento, visita.municipio, visita.address, visita.usuario, visita.sede, visita.estado]
        ws.append(row)

    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            cell.border = thin_border

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=reporte_inspecciones.xlsx'

    wb.save(response)

    return response

@login_required
def otra_foto(request, visita_id):

    if request.user.usuario.userType != 'Correspondencia' and request.user.usuario.userType != 'Admin':
        return redirect('sedes')
    
    if request.method == 'POST':
        form = FotoForm(request.POST, request.FILES)
        if form.is_valid():
            foto = form.save(commit=False)
            foto.visita_id = visita_id

            last_foto = Foto.objects.filter(visita_id=visita_id).order_by('-number').first()
            if last_foto:
                foto.number = str(int(last_foto.number) + 1)
            else:
                foto.number = '1'

            foto.save()
            return redirect('visita_detail', visita_id=visita_id)
    else:
        form = FotoForm()

    return render(request, 'add_foto.html', {
        'form': form
        })
